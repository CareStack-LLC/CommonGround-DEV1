"""
Document Text Extractor for ARIA Agreement Builder.

Extracts text from uploaded PDF and DOCX files so ARIA can use
existing custody/parenting agreements as a conversation starting point.

Uses pymupdf (fitz) for PDFs and python-docx for Word documents —
both already in requirements.txt and used elsewhere in the codebase.
"""

import io
from typing import Tuple

import fitz  # pymupdf
from docx import Document as DocxDocument


# Max chars to return (fits comfortably in GPT-4-Turbo 128K context
# alongside system prompt + conversation history)
MAX_TEXT_LENGTH = 30_000

# If a multi-page PDF yields fewer chars than this, it's likely scanned images
IMAGE_PDF_THRESHOLD = 100


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, int]:
    """
    Extract text from a PDF file.

    Args:
        file_content: Raw PDF bytes

    Returns:
        Tuple of (extracted_text, page_count)

    Raises:
        ValueError: If PDF is image-only (scanned) with no extractable text
        Exception: If PDF is corrupt or unreadable
    """
    doc = fitz.open(stream=file_content, filetype="pdf")
    page_count = len(doc)
    pages_text = []

    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages_text.append(f"--- Page {page_num + 1} ---\n{text}")

    doc.close()

    full_text = "\n\n".join(pages_text)

    # Detect image-only PDFs
    if page_count > 0 and len(full_text) < IMAGE_PDF_THRESHOLD:
        raise ValueError(
            "This PDF appears to contain scanned images rather than text. "
            "ARIA can only read text-based documents. Please upload a "
            "text-based PDF or Word document instead."
        )

    # Truncate if too long
    if len(full_text) > MAX_TEXT_LENGTH:
        full_text = full_text[:MAX_TEXT_LENGTH] + (
            f"\n\n[Document truncated — showing first ~{MAX_TEXT_LENGTH:,} characters]"
        )

    return full_text, page_count


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX (Word) file.

    Extracts both paragraph text and table content.

    Args:
        file_content: Raw DOCX bytes

    Returns:
        Extracted text string

    Raises:
        Exception: If file is corrupt or unreadable
    """
    doc = DocxDocument(io.BytesIO(file_content))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    full_text = "\n\n".join(parts)

    # Truncate if too long
    if len(full_text) > MAX_TEXT_LENGTH:
        full_text = full_text[:MAX_TEXT_LENGTH] + (
            f"\n\n[Document truncated — showing first ~{MAX_TEXT_LENGTH:,} characters]"
        )

    return full_text


def extract_text(file_content: bytes, content_type: str, filename: str) -> Tuple[str, dict]:
    """
    Extract text from a file based on its content type.

    Args:
        file_content: Raw file bytes
        content_type: MIME type of the file
        filename: Original filename

    Returns:
        Tuple of (extracted_text, metadata_dict)
        metadata_dict includes: {page_count, text_length, extraction_method}

    Raises:
        ValueError: For unsupported file types or image-only PDFs
    """
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        text, page_count = extract_text_from_pdf(file_content)
        return text, {
            "page_count": page_count,
            "text_length": len(text),
            "extraction_method": "pymupdf",
        }

    elif (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    ):
        text = extract_text_from_docx(file_content)
        return text, {
            "text_length": len(text),
            "extraction_method": "python-docx",
        }

    else:
        raise ValueError(
            f"Unsupported file type: {content_type}. "
            "Please upload a PDF (.pdf) or Word document (.docx)."
        )
